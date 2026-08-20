"""
Phase 4 · brief.py —— 总裁办简报（Action-triggered RAG）运行期全部逻辑

数据流（两条通道，只在渲染层汇合）：

  数字通道 ── 引擎读数 → 定值字符串 + 数字白名单 ─────────────┐
                                                              ├→ 出口对账 → 渲染
  检索通道 ── trigger_key → 查表取向量 → 余弦 top-k → LLM ────┘

🔴 红线：LLM 幻觉够不到财务数字。
   · 送进 prompt 的只有【定性标签】（裁决态、档位、名次），没有任何数值
   · LLM 只返回措辞与占位符 {slot}，数值由 python 在渲染层填入
   · 出口对账：扫出的每个数字必须在白名单里，且禁止口算式表达（"腰斩""翻倍"）
   · 不过关的句子丢弃并回退模板句，不是给个警告了事

运行期不装嵌入模型、不调嵌入 API —— 向量在 build_corpus.py 里离线算好随仓库走。

冒烟：python brief.py
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field, asdict
from pathlib import Path

import numpy as np

import triggers as TG

try:
    import copy_cn as T                      # 一切人话仍归 copy_cn（宪章 §6）
except ImportError:                          # 允许脱离仓库单独冒烟
    T = None


def _T(name, default):
    return getattr(T, name, default) if T else default

ROOT = Path(__file__).resolve().parent
CHUNKS_PATH = ROOT / "corpus" / "chunks.jsonl"
VECS_PATH = ROOT / "corpus" / "vecs.npz"
CACHE_DIR = ROOT / "cache"

# 供应商与检索参数统一从 config 读（换厂商只改 config.py 三行，本文件不动）
try:
    import config as C
    TOP_K = getattr(C, "RAG_TOP_K", 3)
    BASE_URL = getattr(C, "LLM_BASE_URL", "https://open.bigmodel.cn/api/paas/v4")
    MODEL = getattr(C, "LLM_MODEL", "glm-4-flash")
    TEMPERATURE = getattr(C, "LLM_TEMPERATURE", 0.6)
    MAX_TOKENS = getattr(C, "LLM_MAX_TOKENS", 2000)
    TIMEOUT_S = getattr(C, "LLM_TIMEOUT_S", 60)
except ImportError:                       # 允许脱离仓库单独冒烟
    TOP_K, TEMPERATURE, MAX_TOKENS, TIMEOUT_S = 3, 0.6, 2000, 60
    BASE_URL, MODEL = "https://open.bigmodel.cn/api/paas/v4", "glm-4-flash"


# ══════════════════════════════════════════════════════════════
# 一、数字通道：读数 → 定值字符串 + 白名单
# ══════════════════════════════════════════════════════════════
def pct(x, digits=1):
    """百分比定值渲染。全流程只认这一个形态 —— 白名单按【渲染后的字符串】建，
    否则 LLM 复述 '-5.2%' 而白名单里是 '-0.0517'，对账会全部误杀。"""
    if x is None or (isinstance(x, float) and x != x):
        return "—"
    return f"{x * 100:+.{digits}f}%"


def pctu(x, digits=0):
    """无符号百分比 —— 用于【水平量】（份额），与带符号的【变动量】区分开。
    份额是"占多少"，不是"变了多少"，写成 +38% 会误导，也会让对账误杀。"""
    if x is None or (isinstance(x, float) and x != x):
        return "—"
    return f"{x * 100:.{digits}f}%"


def num(x, digits=2):
    if x is None or (isinstance(x, float) and x != x):
        return "—"
    return f"{x:.{digits}f}"


def _band(v, cuts, names):
    """连续值 → 档位标签。送给 LLM 的是这个词，不是数值。"""
    for c, n in zip(cuts, names):
        if v <= c:
            return n
    return names[-1]


@dataclass
class Readout:
    """引擎读数包。全部字段由 app.py 从已算好的结果填入，brief.py 不做任何计算。"""
    city: str = ""
    city_cn: str = ""
    quad: str = ""
    quad_cn: str = ""
    # 旋钮
    price_pct: float = 0.0
    eco: float = 0.0
    ally: bool = False
    shock_pct: float = 0.0
    ruler_cn: str = ""
    # 图 A · 你自己
    roe_base: float | None = None
    roe_end: float | None = None
    spread_end: float | None = None
    beta_used: float | None = None
    gamma_used: float | None = None
    net_margin: float | None = None
    asset_turnover: float | None = None
    equity_multiplier: float | None = None
    # 图 B/C · 博弈
    verdict_state: str = "CREATE"
    verdict_sentence: str = ""
    share: float | None = None
    share_rank: str = "—"
    spread_rank: str = "—"
    spread_game: float | None = None
    a_value: float | None = None
    in_alliance: bool = False
    competition_cn: str = ""

    # ── 定性标签：唯一允许进入 prompt 的"程度"表达 ──────────────
    def labels(self) -> dict:
        return {
            "城市": self.city_cn or self.city,
            "象限": self.quad_cn or self.quad,
            "裁决态": _T("STATE_CN", {}).get(self.verdict_state, self.verdict_state),
            "定价动作": _band(self.price_pct, [-20, -8, -1, 1, 8],
                              ["大幅下调", "明显下调", "小幅下调", "维持不变",
                               "小幅上调", "明显上调"]),
            "生态投资": _band(self.eco, [0.001, 0.3, 0.7],
                              ["未投入", "少量投入", "中等投入", "大力投入"]),
            "换电联盟": "已加入" if self.ally else "未加入",
            "原材料冲击": _band(self.shock_pct, [-1, 1, 20],
                                ["价格下行", "基本平稳", "温和上涨", "大幅上涨"]),
            "账面回报": ("为正" if (self.roe_end or 0) > 0 else "为负"),
            "价值利差": ("为正" if (self.spread_end or 0) > 0 else "为负"),
            "份额名次": str(self.share_rank),
            "价值名次": str(self.spread_rank),
            "评判指标": self.ruler_cn,
            "竞争类型": self.competition_cn,
        }

    # ── 定值字符串：渲染层填模板用，【不】进 prompt ──────────────
    def values(self) -> dict:
        return {
            "city": self.city_cn or self.city,
            "quad": self.quad_cn or self.quad,
            "price_pct": f"{self.price_pct:+.0f}%",
            "eco": num(self.eco),
            "ally": "已加入" if self.ally else "未加入",
            "shock_pct": f"{self.shock_pct:+.0f}%",
            "ruler": self.ruler_cn,
            "roe_base": pct(self.roe_base),
            "roe_end": pct(self.roe_end),
            "spread_end": pct(self.spread_end),
            "beta_used": num(self.beta_used),
            "gamma_used": num(self.gamma_used),
            "net_margin": pct(self.net_margin),
            "asset_turnover": num(self.asset_turnover),
            "equity_multiplier": num(self.equity_multiplier),
            "share": pctu(self.share),
            "share_rank": str(self.share_rank),
            "spread_rank": str(self.spread_rank),
            "spread_game": pct(self.spread_game),
            "a_value": num(self.a_value),
            "alliance": "在盟" if self.in_alliance else "未在盟",
            "competition": self.competition_cn,
            "verdict_cn": _T("STATE_CN", {}).get(self.verdict_state,
                                                 self.verdict_state),
        }

    def whitelist(self) -> set:
        """允许出现在简报里的数字（按渲染后的形态）。"""
        out = set()
        for v in self.values().values():
            for n in _NUM_RE.findall(str(v)):
                out.add(n)
                out.add(n.lstrip("+"))        # "+5.1%" 亦允许写作 "5.1%"
        return out

    def trigger_key(self) -> str:
        return TG.trigger_key(self.quad, self.verdict_state)


# ══════════════════════════════════════════════════════════════
# 二、检索通道：查表 → 点积 → top-k（无嵌入模型、无 API）
# ══════════════════════════════════════════════════════════════
_STORE = {}


def _load():
    if _STORE:
        return _STORE
    if not (CHUNKS_PATH.exists() and VECS_PATH.exists()):
        _STORE.update(ok=False)
        return _STORE
    chunks = [json.loads(l) for l in CHUNKS_PATH.open(encoding="utf-8") if l.strip()]
    Z = np.load(VECS_PATH, allow_pickle=False)
    _STORE.update(
        ok=True, chunks=chunks,
        cvecs=Z["corpus_vecs"],
        qpos={k: i for i, k in enumerate(Z["query_keys"].tolist())},
        qvecs=Z["query_vecs"],
        rows_by_view={v: [i for i, c in enumerate(chunks) if c["view"] == v]
                      for v in TG.VIEWS},
    )
    return _STORE


def retrieve(city: str, quad: str, state: str, k: int = TOP_K) -> dict:
    """两个视角各检索一次，按 view 分池 → 两边取回的材料天然不重复。

    检索判据与 build_corpus.diagnose 完全一致（结构化元数据优先、语义补位）：
      policy   : 先锚定 1 条该城专属材料，再由相似度从「该城 + 全国性」池补足
      strategy : 先锚定「你的裁决态」「你的象限」两条，再由相似度补第三条
    纯相似度会让泛化性强的切片成为枢纽，把专属材料挤掉（v1 实测 17/20 格被同一条占据）。
    """
    S = _load()
    if not S.get("ok"):
        return {v: [] for v in TG.VIEWS}
    C, V = S["chunks"], S["cvecs"]

    def _sim_order(rows, qvec):
        sims = {i: float(V[i] @ qvec) for i in rows}
        return sorted(rows, key=lambda i: -sims[i])

    def _fill(out, ordered):
        for i in ordered:
            if len(out) >= k:
                break
            if i not in out:
                out.append(i)
        return out[:k]

    out = {}

    pos = S["qpos"].get(f"policy::{TG.policy_key(city, state)}")
    if pos is None:
        out["policy"] = []
    else:
        qv = S["qvecs"][pos]
        rows = [i for i, c in enumerate(C)
                if c["view"] == "policy" and c.get("city") in (city, None)]
        own = [i for i in rows if C[i].get("city") == city]
        picked = [_sim_order(own, qv)[0]] if own else []
        out["policy"] = [C[i] for i in _fill(picked, _sim_order(rows, qv))]

    pos = S["qpos"].get(f"strategy::{TG.strategy_key(quad, state)}")
    if pos is None:
        out["strategy"] = []
    else:
        qv = S["qvecs"][pos]
        rows = [i for i, c in enumerate(C) if c["view"] == "strategy"]
        picked = []
        for key, val in (("state", state), ("quad", quad)):
            hit = next((i for i in rows if C[i].get(key) == val), None)
            if hit is not None and hit not in picked:
                picked.append(hit)
        out["strategy"] = [C[i] for i in _fill(picked, _sim_order(rows, qv))]

    return out


# ══════════════════════════════════════════════════════════════
# 三、生成层：一次调用，只出措辞
# ══════════════════════════════════════════════════════════════
SLOTS = ("summary", "combined", "policy_view", "compete_view", "conclusion")

_SYSTEM = """你是一位面向管理层的战略分析写作助手。你的唯一任务是【组织措辞】。

分析主体（务必看清）：
本简报分析的是【一家车企】——即"本局您操控的车企"。
它的"所在地"与"战略象限"只是这家企业的两个属性：所在地说明它在哪里建厂、
继承了哪份区域禀赋；象限说明它做什么档次、什么动力路线的车。

参考材料的身份：下方政策类材料描述的是【这家企业所在地的产业环境】，
是它面对的背景条件，**不是分析对象**。读到"某市产业链如何"时，
你要写的是"该企业受当地配套条件影响，应如何应对"，而不是复述该市的产业状况。

正误示例：
✗ 西安新能源汽车产业面临零部件配套不足，政策鼓励产业链完善。
✓ 该企业所在地的零部件配套尚不完整，核心部件需外购，抬高了其成本刚性；
  在当前定价动作下，这一约束会被进一步放大。
✗ 西安应重视价值创造。
✓ 该企业应把目标从份额转向价值创造。

绝对禁止（违反则整段作废）：
1. 不得写出任何具体数值：百分比、金额、倍数、名次数字，一律不写。
2. 不得做任何计算或比较换算，包括"腰斩""翻倍""高出三成""接近一半"这类
   没有阿拉伯数字但实质在做算术的表达。
3. 不得编造材料中没有的政策、企业或事实。

你会收到：
· 局面标签：这一局的定性描述（不含数值）
· 参考材料：政策与战略材料摘编，可引用其中的观点与做法

只输出一个 JSON 对象，不要 markdown 代码块，不要任何解释。字段：
{
  "summary":      "本轮摘要，120-160字，说清做了什么、落到什么局面、建议方向",
  "combined":     "综合分析，120-160字，把几个动作合起来看是什么局面",
  "policy_view":  "政策与区位视角，200-260字，结合参考材料",
  "compete_view": "竞争与生态视角，200-260字，结合参考材料",
  "conclusion":   "总结，150-200字，收口并给出下一步方向"
}

写作要求：
· 主语一律用「该企业」「这家企业」，或直接省略主语。
· 城市名只能出现在"依托当地…""受本地配套限制…"这类**区位状语**里，
  绝不可作为句子的主语。写出"西安应该…""合肥需要…"即为错误。
· 若参考材料为空，就只依据局面标签写，不要编造政策内容，也不要说"材料未提供"。
· 不要照搬参考材料的标题或其中的英文代号、企业名清单；把材料的观点用自己的话写进句子。
· 禁止空话套话：不得出现"提升核心竞争力""实现可持续发展""推动高质量发展""赋能"
  这类不含具体动作的词组。每一句建议都要落到可执行的动作上
  （例如"淘汰回报最低的入门车型"而非"优化产品结构"）。

再次强调两点：
一、分析对象是【一家车企】，不是城市；主语用「该企业」。
二、直接以 { 开头、以 } 结尾输出，前后不要有任何说明文字或代码围栏。"""


def _prompt(labels: dict, docs: dict) -> str:
    # 先用一句话把主体钉死，再列其余标签 —— 否则模型会把城市误当成分析对象（实测）
    subject = (f"分析对象：{_T('BRIEF_SUBJECT', '本局您操控的车企')}"
               f"——一家定位为「{labels.get('象限', '')}」的车企。\n"
               f"· 所在地（区位背景，非分析对象）：{labels.get('城市', '')}")
    rest = {k: v for k, v in labels.items() if k not in ("城市", "象限")}
    lab = subject + "\n" + "\n".join(f"· {k}：{v}" for k, v in rest.items())
    def _clean_title(t: str) -> str:
        """剥掉切片标题里的内部键名前缀。

        strategy.md 的标题形如「WIN_ALONE：从份额最大化转向价值最大化」，
        直接丢进 prompt，模型会照抄"借鉴 WIN_ALONE 的理念"——内部键名漏进正文（实测）。
        这里只保留冒号后的人话部分。
        """
        head, sep, tail = t.partition("：")
        if sep and (head in TG.STATES or re.match(r"^Q[1-4]\b", head)):
            return tail.strip() or head
        return t

    def _fmt(items):
        return "\n".join(f"[{_clean_title(c['title'])}]\n{c['body']}"
                          for c in items) or "（无）"
    return (f"局面标签：\n{lab}\n\n"
            f"参考材料 · 政策类：\n{_fmt(docs.get('policy', []))}\n\n"
            f"参考材料 · 战略类：\n{_fmt(docs.get('strategy', []))}\n\n"
            # 小模型对开头与结尾最敏感，中段指令易被稀释 → 首尾各钉一次主体
            f"请以【该企业】为主语撰写。上方政策材料是它所在地的背景条件，"
            f"不要把城市或当地产业写成主语。")


def _extract_json(txt: str) -> dict:
    """从模型输出里挖出 JSON 对象。

    实测（GLM-4-Flash）：模型常在 JSON 前后裹一层说明文字或 markdown 围栏，
    只剥围栏是不够的。故三级兜底：
      ① 直接解析 → ② 剥 ``` 围栏后解析 → ③ 取最外层 {...} 子串解析
    """
    txt = (txt or "").strip()
    try:
        return json.loads(txt)
    except json.JSONDecodeError:
        pass
    stripped = re.sub(r"```(?:json)?|```", "", txt).strip()
    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        pass
    i, j = stripped.find("{"), stripped.rfind("}")
    if i != -1 and j > i:
        return json.loads(stripped[i:j + 1])          # 失败则抛给上层
    raise json.JSONDecodeError("no json object", stripped or "<empty>", 0)


def generate(labels: dict, docs: dict, api_key: str) -> tuple[dict, str]:
    """返回 (slots, error)。任何失败都不抛异常 —— 上层据 error 走降级路径。"""
    try:
        from openai import OpenAI
    except ImportError:
        return {}, "缺少依赖：pip install openai"

    cli = OpenAI(api_key=api_key, base_url=BASE_URL, timeout=TIMEOUT_S)
    msgs = [{"role": "system", "content": _SYSTEM},
            {"role": "user", "content": _prompt(labels, docs)}]
    kw = dict(model=MODEL, temperature=TEMPERATURE, max_tokens=MAX_TOKENS,
              messages=msgs)

    txt = ""
    try:
        # 先试 JSON 模式（智谱 GLM-4 系列支持）；不支持的模型会报错，再退回普通模式。
        try:
            r = cli.chat.completions.create(
                response_format={"type": "json_object"}, **kw)
        except Exception:                                     # noqa: BLE001
            r = cli.chat.completions.create(**kw)
        txt = (r.choices[0].message.content or "").strip()
        data = _extract_json(txt)
        return {k: str(data.get(k, "")).strip() for k in SLOTS}, ""
    except json.JSONDecodeError:
        # 把原始输出的开头带回去 —— 否则"不是合法 JSON"这句话无从排查
        head = txt[:160].replace("\n", " ") if txt else "<空响应>"
        return {}, f"模型返回的不是合法 JSON；原始输出开头：{head}"
    except Exception as e:                                    # noqa: BLE001
        return {}, f"{type(e).__name__}: {e}"


# ══════════════════════════════════════════════════════════════
# 四、出口对账：红线的执行机构
# ══════════════════════════════════════════════════════════════
_NUM_RE = re.compile(r"[-+]?\d+(?:\.\d+)?%?")
# 没有阿拉伯数字、但实质在做算术的表达。命中即判该句不合格。
_ARITH_WORDS = ("腰斩", "翻倍", "翻番", "成倍", "减半", "倍增", "折半",
                "三成", "四成", "五成", "六成", "七成", "八成", "九成",
                "过半", "近半", "一半", "三分之", "四分之", "百分之",
                "高出", "低出", "多出", "少了", "增加了", "减少了")
_SENT_SPLIT = re.compile(r"(?<=[。！？；\n])")


def reconcile(text: str, allowed: set) -> tuple[str, list]:
    """逐句审：句中每个数字必须在白名单；出现口算词即判不合格。
    不合格的句子【丢弃】，不是标注警告 —— 宁可少一句，不可脏一个数。"""
    kept, dropped = [], []
    for sent in _SENT_SPLIT.split(text):
        s = sent.strip()
        if not s:
            continue
        bad = [n for n in _NUM_RE.findall(s) if n not in allowed]
        word = next((w for w in _ARITH_WORDS if w in s), None)
        if bad:
            dropped.append((s, f"数字不在白名单：{'、'.join(bad)}"))
        elif word:
            dropped.append((s, f"口算式表达：{word}"))
        else:
            kept.append(s)
    return "".join(kept), dropped


def clean_slots(slots: dict, allowed: set, fallback: dict) -> tuple[dict, list]:
    out, drops = {}, []
    for k in SLOTS:
        txt, d = reconcile(slots.get(k, ""), allowed)
        drops += [(k, s, why) for s, why in d]
        out[k] = txt if len(txt) >= 20 else fallback.get(k, "")   # 整段被清空 → 回退
    return out, drops


# ══════════════════════════════════════════════════════════════
# 五、组装：§二 全部由模板 + 定值字符串生成（零 LLM、零风险）
# ══════════════════════════════════════════════════════════════
def action_rows(r: Readout) -> list[tuple]:
    """§二 的读数表：动作 → 本轮取值 → 对账面的影响 → 对价值/博弈的影响。
    每一格都由引擎读数决定，方向是确定的，不需要模型推断。"""
    v = r.values()
    lab = r.labels()
    return [
        ("定价", v["price_pct"],
         f"经价格弹性 β={v['beta_used']} 传导到销量；账面回报 {v['roe_end']}",
         f"价值利差 {v['spread_end']}；份额 {v['share']}、名次 {v['share_rank']}"),
        ("生态投资", v["eco"],
         f"抬升非价格吸引力至 {v['a_value']}，并作为需求侧位移驱动账面",
         f"价值名次 {v['spread_rank']}；{lab['生态投资']}"),
        ("换电联盟", v["ally"],
         "不直接改动账面成本结构",
         f"跨象限竞合中{v['alliance']}"),
        ("关键原材料价格冲击", v["shock_pct"],
         f"经成本传导 γ={v['gamma_used']} 打到单位成本，按半衰期衰减",
         "外生环境，非你的决策"),
    ]


def dupont_rows(r: Readout) -> list[tuple]:
    v = r.values()
    return [("净利率", v["net_margin"], "每一块钱收入剩下多少"),
            ("资产周转", v["asset_turnover"], "一块钱资产转出多少收入"),
            ("权益乘数", v["equity_multiplier"], "自有资金撬动了多少总资产")]


def table_leads(r: Readout) -> dict:
    """两张表的引导叙述。纯模板 + 定性标签，零 LLM 参与、零编造风险。
    报告惯例：先叙述、再引「（见表 N）」、最后出表 —— 表格是论据，不替代论述。"""
    lab = r.labels()
    ref = _T("BRIEF_TABLE_REF", {"action": "（见表 1）", "dupont": "（见表 2）"})
    lead_a = _T("BRIEF_LEAD_ACTION", "")
    lead_d = _T("BRIEF_LEAD_DUPONT", "")
    return {
        "action": lead_a.format(price_lab=lab["定价动作"], eco_lab=lab["生态投资"],
                                ally_lab=lab["换电联盟"], shock_lab=lab["原材料冲击"],
                                ref=ref["action"]) if lead_a else "",
        "dupont": lead_d.format(ref=ref["dupont"]) if lead_d else "",
    }


def fallback_slots(r: Readout) -> dict:
    """降级模板：无 key、调用失败、或某段被对账清空时使用。纯模板，不含推断。"""
    lab = r.labels()
    v = r.values()
    return {
        "summary": (f"本轮在{v['city']}以「{v['quad']}」定位推演："
                    f"定价{lab['定价动作']}，生态投资{lab['生态投资']}，"
                    f"换电联盟{lab['换电联盟']}，外部原材料价格{lab['原材料冲击']}。"
                    f"账面回报{lab['账面回报']}，价值利差{lab['价值利差']}，"
                    f"份额名次 {v['share_rank']}、价值名次 {v['spread_rank']}。"),
        # 降级文案刻意【不复述读数】—— 第一段已把数字讲完，这里只讲机理与下一步，
        # 否则屏上两段会说同一件事。
        "combined": (f"定价经价格弹性改写销量与单价，生态投资经非价格吸引力改写需求，"
                     f"外部原材料价格则经成本传导改写单位成本；三条通道先汇入损益，"
                     f"再由损益分别进入价值与博弈两张记分表。"
                     f"当前竞争类型为{v['competition']}，"
                     f"这决定了同象限对手对你这一步的还手力度。"),
        "policy_view": "（本段需要调用模型生成，当前使用引擎读数版简报。）",
        "compete_view": "（本段需要调用模型生成，当前使用引擎读数版简报。）",
        "conclusion": (f"下一步可从三处着手：一是换一项评判指标复看同一局，"
                       f"确认结论在不同尺度下是否翻转；二是把定价与生态投资反向对调，"
                       f"看非价格手段能否替代降价；三是切换城市或象限，"
                       f"检验当前结果是禀赋决定的，还是定位决定的。"),
    }


def build(readout: Readout, api_key: str = "") -> dict:
    """总装。返回一份 report dict —— 屏上与 docx 共用同一份，绝不二次调用模型。"""
    tk = readout.trigger_key()
    docs = retrieve(readout.city, readout.quad, readout.verdict_state)
    fb = fallback_slots(readout)

    if api_key:
        slots, err = generate(readout.labels(), docs, api_key)
    else:
        slots, err = {}, "未填写 API Key"

    if slots:
        slots, drops = clean_slots(slots, readout.whitelist(), fb)
        mode = "live"
    else:
        slots, drops, mode = fb, [], "fallback"

    cites = [{"title": c["title"], "sources": c["sources"], "view": c["view"]}
             for v in TG.VIEWS for c in docs.get(v, [])]
    return dict(trigger_key=tk, mode=mode, error=err, slots=slots,
                values=readout.values(), labels=readout.labels(),
                leads=table_leads(readout),
                action_rows=action_rows(readout), dupont_rows=dupont_rows(readout),
                cites=cites, dropped=drops,
                verdict=readout.verdict_sentence, readout=asdict(readout))


# ══════════════════════════════════════════════════════════════
# 六、渲染 · 屏上（§一 + §四，约 500 字）
# ══════════════════════════════════════════════════════════════
def to_markdown(rep: dict) -> str:
    """屏上简报 = 两段，无章节大标题（大标题只在导出的完整版里用）。

      第一段：动作与读数复述 —— 全部由 python 填引擎定值字符串，数字最密、零自由度
      第二段：影响与策略 —— LLM 组织措辞，不碰任何数字

    这样"报数字"和"讲道理"各占一段，读者一眼分得清哪句是算出来的、哪句是解释。
    """
    s = rep["slots"]
    recap = _T("BRIEF_SCREEN_RECAP", "")
    para1 = recap.format(**rep["values"]) if recap else s["summary"]
    para2 = "".join(x for x in (s.get("combined", ""), s.get("conclusion", "")) if x)
    parts = [para1, para2,
             f"_{_T('BRIEF_TAIL_HINT', '完整版见导出文档。')}_"]
    note = _T("DUAL_BASIS_NOTE", "")
    if note:
        parts.append(f":gray[{note}]")            # 两套口径的诚实标注
    if rep["cites"]:
        srcs = []
        for c in rep["cites"]:
            for x in c["sources"]:
                if x not in srcs:
                    srcs.append(x)
        title = _T("BRIEF_CITE_TITLE", "引用材料")
        parts.append(f"**{title}**\n\n" + "\n".join(f"- {x}" for x in srcs))
    return "\n\n".join(parts)


# ══════════════════════════════════════════════════════════════
# 七、渲染 · docx（两页，宋体 + Calibri/Arial）
# ══════════════════════════════════════════════════════════════
def _font(run, size, bold=False, latin="Calibri", east="宋体"):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = latin
    # python-docx 的 font.name 只写 ascii/hAnsi，从不碰 eastAsia →
    # 中文字符会 fallback 到 Word 默认字体。必须显式设 w:eastAsia。
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)


def _para(doc, text, size=12, bold=False, level=0, latin="Calibri"):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 0:                                    # 正文
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Pt(size * 2)           # 首行缩进 2 字符
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.35
    else:                                             # 标题
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        before, after, exact = {1: (12, 6, 24), 2: (6, 4, 20), 3: (3, 2, 18)}[level]
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(exact)
    _font(p.add_run(text), size, bold, latin=latin)
    return p


def _caption(doc, text):
    """表格题注，置于表格【上方】（报告惯例：表题在上、图题在下）。"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before, pf.space_after = Pt(6), Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _font(p.add_run(text), 10.5, bold=True, latin="Arial")
    return p


def _cell(cell, text, size=10.5, bold=False):
    """单元格内容：水平居中 + 垂直居中。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = para.paragraph_format.space_after = None
    _font(para.add_run(str(text)), size, bold=bold, latin="Arial")


def _table(doc, headers, rows, widths=None, caption=None):
    from docx.shared import Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    if caption:
        _caption(doc, caption)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            _cell(cells[i], val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def to_docx(rep: dict, path: str | Path) -> Path:
    """两页简报。与屏上同源同字：吃的是同一份 rep，不再调用模型。"""
    from docx import Document
    from docx.shared import Pt

    v, s = rep["values"], rep["slots"]
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(56)
        sec.left_margin = sec.right_margin = Pt(56)

    H1 = _T("BRIEF_H1", {"summary": "一、本轮摘要", "action": "二、本局行动分析",
                         "strategy": "三、策略分析", "conclusion": "四、总结"})
    H2 = _T("BRIEF_H2", {"action_rows": "2.1 逐项动作与影响",
                         "dupont": "2.2 账面影响（杜邦分析）",
                         "combined": "2.3 综合分析",
                         "policy_view": "3.1 政策与区位视角",
                         "compete_view": "3.2 竞争与生态视角"})
    CAP = _T("BRIEF_TABLE_CAPTION", {"action": "表 1　本轮动作及其影响",
                                     "dupont": "表 2　账面结构的三因子拆解"})
    leads = rep.get("leads", {})

    _para(doc, _T("BRIEF_DOC_TITLE", "沙盘模拟下新能源汽车企业动态决策的商业分析"),
          16, True, level=1, latin="Arial")
    _para(doc, _T("BRIEF_DOC_SUBTITLE",
                  "城市选址：{city}　|　企业特征：{quad}　|　评判指标：{ruler}")
          .format(city=v["city"], quad=v["quad"], ruler=v["ruler"]),
          10.5, level=3, latin="Arial")

    _para(doc, H1["summary"], 16, True, level=1, latin="Arial")
    _para(doc, s["summary"])

    _para(doc, H1["action"], 16, True, level=1, latin="Arial")
    _para(doc, H2["action_rows"], 14, True, level=2, latin="Arial")
    if leads.get("action"):
        _para(doc, leads["action"])                 # 先叙述，再出表
    _table(doc, _T("BRIEF_TABLE_ACTION",
                   ["动作", "本轮取值", "对账面的影响", "对价值与博弈的影响"]),
           rep["action_rows"], widths=[2.4, 2.0, 5.4, 5.4], caption=CAP["action"])

    _para(doc, H2["dupont"], 14, True, level=2, latin="Arial")
    if leads.get("dupont"):
        _para(doc, leads["dupont"])
    _table(doc, _T("BRIEF_TABLE_DUPONT", ["因子", "取值", "术语解释"]),
           rep["dupont_rows"], widths=[2.6, 2.4, 10.2], caption=CAP["dupont"])

    _para(doc, H2["combined"], 14, True, level=2, latin="Arial")
    _para(doc, s["combined"])

    _para(doc, H1["strategy"], 16, True, level=1, latin="Arial")
    _para(doc, H2["policy_view"], 14, True, level=2, latin="Arial")
    _para(doc, s["policy_view"])
    _para(doc, H2["compete_view"], 14, True, level=2, latin="Arial")
    _para(doc, s["compete_view"])

    _para(doc, H1["conclusion"], 16, True, level=1, latin="Arial")
    _para(doc, s["conclusion"])

    # ── 引用材料：另起一页 ──
    if rep["cites"]:
        doc.add_page_break()
        _para(doc, _T("BRIEF_CITE_TITLE", "引用材料"), 16, True, level=1, latin="Arial")
        seen = []
        for c in rep["cites"]:
            for x in c["sources"]:
                if x not in seen:
                    seen.append(x)
        for i, x in enumerate(seen, 1):
            _para(doc, f"[{i}] {x}", 9, level=3, latin="Arial")

    if _T("DUAL_BASIS_NOTE", ""):
        _para(doc, _T("DUAL_BASIS_NOTE", ""), 9, level=3, latin="Arial")
    _para(doc, _T("BRIEF_DISCLAIMER",
                  "本简报中的全部数值均由确定性引擎计算并直接填入，"
                  "语言模型仅负责组织措辞、不参与任何计算。"
                  "本模型为简化教学与娱乐用途，非真实预测。"),
          9, level=3, latin="Arial")

    path = Path(path)
    doc.save(path)
    return path


# ══════════════════════════════════════════════════════════════
# 八、演示缓存
# ══════════════════════════════════════════════════════════════
def cache_path(trigger_key: str) -> Path:
    return CACHE_DIR / f"{trigger_key.replace('|', '_')}.json"


def load_cache(trigger_key: str) -> dict | None:
    p = cache_path(trigger_key)
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:                                     # noqa: BLE001
            return None
    return None


def save_cache(rep: dict) -> Path:
    CACHE_DIR.mkdir(exist_ok=True)
    p = cache_path(rep["trigger_key"])
    p.write_text(json.dumps(rep, ensure_ascii=False, indent=1), encoding="utf-8")
    return p


# ══════════════════════════════════════════════════════════════
# 冒烟：python brief.py
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    r = Readout(city="Xian", city_cn="西安", quad="Q4", quad_cn="极致性价比",
                price_pct=-12, eco=0.2, ally=False, shock_pct=20, ruler_cn="价值利差",
                roe_base=0.051, roe_end=-0.081, spread_end=-0.052,
                beta_used=-1.35, gamma_used=0.62,
                net_margin=-0.021, asset_turnover=0.88, equity_multiplier=2.4,
                verdict_state="WIN_ALONE", verdict_sentence="份额第一，价值垫底。",
                share=0.38, share_rank="1", spread_rank="5", spread_game=-0.052,
                a_value=1.12, in_alliance=False, competition_cn="价格主导")

    print("trigger_key:", r.trigger_key())
    print("白名单:", sorted(r.whitelist())[:12], "…")

    docs = retrieve(r.city, r.quad, r.verdict_state)
    for view in TG.VIEWS:
        got = [c["title"][:24] for c in docs[view]]
        print(f"检索 {view}: {got or '（未构建语料，跑 build_corpus.py）'}")

    # 对账自检：白名单内数字放行，白名单外与口算词丢弃
    ok, dropped = reconcile(
        "价值利差为 -5.2%，份额 38%。利差较上轮腰斩。回报率高达 -9.9%。",
        r.whitelist())
    print("\n对账保留:", ok)
    for s, why in dropped:
        print("  丢弃:", s, "→", why)

    rep = build(r, api_key="")           # 无 key → 走降级模板
    print("\n模式:", rep["mode"], "|", rep["error"])
    print("\n" + to_markdown(rep))

    try:
        out = to_docx(rep, ROOT / "brief_smoke.docx")
        print(f"\ndocx 已写出：{out}")
    except ImportError:
        print("\n（未装 python-docx，跳过 docx 冒烟：pip install python-docx）")
